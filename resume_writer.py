"""
Resume writer — generates a clean, ATS-friendly DOCX from a resume dict.

Uses python-docx to produce a professional Word document with:
  - Name as Heading 1
  - Contact information line
  - Horizontal rule separator
  - Summary, Experience, Education, Skills, Certifications sections
"""

import logging
import os
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# ---------------------------------------------------------------------------
# Type alias
# ---------------------------------------------------------------------------
ResumeDict = Dict[str, Any]


# ---------------------------------------------------------------------------
# Colour constants
# ---------------------------------------------------------------------------
_COLOR_NAME = RGBColor(0x1A, 0x1A, 0x2E)       # Dark navy for name
_COLOR_HEADING = RGBColor(0x16, 0x21, 0x3E)     # Section headings
_COLOR_SUBHEADING = RGBColor(0x0F, 0x3C, 0x78)  # Job titles / degrees
_COLOR_BODY = RGBColor(0x33, 0x33, 0x33)        # Body text
_COLOR_MUTED = RGBColor(0x66, 0x66, 0x66)       # Dates / secondary info
_COLOR_RULE = RGBColor(0x20, 0x60, 0xC0)        # Horizontal rule accent


class ResumeWriter:
    """
    Write a tailored resume dict to a professionally formatted DOCX file.

    Args:
        logger: Optional logger; a default one is created if omitted.
    """

    def __init__(self, logger: Optional[logging.Logger] = None) -> None:
        self.logger = logger or logging.getLogger(__name__)

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def write(self, resume: ResumeDict, output_path: str) -> str:
        """
        Generate a DOCX file at *output_path* from the *resume* dict.

        Args:
            resume:      Structured resume dict (from :class:`ResumeParser`
                         or :class:`TailoringEngine`).
            output_path: Destination file path (must end in ``.docx``).

        Returns:
            The absolute path to the written DOCX file.

        Raises:
            OSError: If the output directory cannot be created or the file
                     cannot be written.
        """
        os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)

        doc = Document()
        self._configure_page(doc)

        # ── Header: name + contact ────────────────────────────────────
        self._add_name(doc, resume.get("name", ""))
        self._add_contact_line(doc, resume)
        self._add_horizontal_rule(doc)

        # ── Summary ───────────────────────────────────────────────────
        summary = resume.get("summary", "").strip()
        if summary:
            self._add_section_heading(doc, "Professional Summary")
            self._add_paragraph(doc, summary)

        # ── Experience ────────────────────────────────────────────────
        experience = resume.get("experience", [])
        if experience:
            self._add_section_heading(doc, "Experience")
            for job in experience:
                self._add_job(doc, job)

        # ── Education ─────────────────────────────────────────────────
        education = resume.get("education", [])
        if education:
            self._add_section_heading(doc, "Education")
            for edu in education:
                self._add_education_entry(doc, edu)

        # ── Skills ────────────────────────────────────────────────────
        skills: List[str] = resume.get("skills", [])
        if skills:
            self._add_section_heading(doc, "Skills")
            self._add_skills(doc, skills)

        # ── Certifications ────────────────────────────────────────────
        certs: List[str] = resume.get("certifications", [])
        if certs:
            self._add_section_heading(doc, "Certifications")
            for cert in certs:
                self._add_bullet(doc, cert)

        doc.save(output_path)
        self.logger.info("Saved tailored resume to: %s", output_path)
        return os.path.abspath(output_path)

    # ------------------------------------------------------------------
    # Page setup
    # ------------------------------------------------------------------

    def _configure_page(self, doc: Document) -> None:
        """Set page margins and default font."""
        section = doc.sections[0]
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

        # Set default font
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)
        font.color.rgb = _COLOR_BODY

    # ------------------------------------------------------------------
    # Name & contact
    # ------------------------------------------------------------------

    def _add_name(self, doc: Document, name: str) -> None:
        """Add the candidate name as a large heading paragraph."""
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(name.upper() if name else "YOUR NAME")
        run.bold = True
        run.font.size = Pt(22)
        run.font.color.rgb = _COLOR_NAME
        run.font.name = "Calibri Light"
        # Remove spacing after
        para.paragraph_format.space_after = Pt(2)

    def _add_contact_line(self, doc: Document, resume: ResumeDict) -> None:
        """Add a single contact line: email  |  phone."""
        parts = []
        if resume.get("email"):
            parts.append(resume["email"])
        if resume.get("phone"):
            parts.append(resume["phone"])

        if not parts:
            return

        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_after = Pt(4)
        run = para.add_run("  |  ".join(parts))
        run.font.size = Pt(10)
        run.font.color.rgb = _COLOR_MUTED

    # ------------------------------------------------------------------
    # Horizontal rule
    # ------------------------------------------------------------------

    def _add_horizontal_rule(self, doc: Document) -> None:
        """Insert a thin horizontal rule paragraph."""
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after = Pt(8)
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "2060C0")
        pBdr.append(bottom)
        pPr.append(pBdr)

    # ------------------------------------------------------------------
    # Section headings
    # ------------------------------------------------------------------

    def _add_section_heading(self, doc: Document, text: str) -> None:
        """Add a section heading in small-caps style."""
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(10)
        para.paragraph_format.space_after = Pt(4)
        run = para.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = _COLOR_HEADING
        run.font.name = "Calibri"

        # Add bottom border to heading paragraph
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "16213E")
        pBdr.append(bottom)
        pPr.append(pBdr)

    # ------------------------------------------------------------------
    # Content blocks
    # ------------------------------------------------------------------

    def _add_paragraph(self, doc: Document, text: str) -> None:
        """Add a regular body paragraph."""
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(6)
        run = para.add_run(text)
        run.font.size = Pt(11)
        run.font.color.rgb = _COLOR_BODY

    def _add_bullet(self, doc: Document, text: str) -> None:
        """Add a bullet-point paragraph."""
        para = doc.add_paragraph(style="List Bullet")
        para.paragraph_format.space_after = Pt(2)
        para.paragraph_format.left_indent = Inches(0.25)
        run = para.add_run(text)
        run.font.size = Pt(11)
        run.font.color.rgb = _COLOR_BODY

    def _add_job(self, doc: Document, job: Dict[str, Any]) -> None:
        """Add a single work experience entry."""
        title = job.get("title", "")
        company = job.get("company", "")
        dates = job.get("dates", "")
        bullets: List[str] = job.get("bullets", [])

        # Job title + company on one line, dates right-aligned
        title_para = doc.add_paragraph()
        title_para.paragraph_format.space_before = Pt(6)
        title_para.paragraph_format.space_after = Pt(1)

        title_run = title_para.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(11)
        title_run.font.color.rgb = _COLOR_SUBHEADING

        if company:
            sep_run = title_para.add_run("  ·  ")
            sep_run.font.color.rgb = _COLOR_MUTED
            sep_run.font.size = Pt(11)

            company_run = title_para.add_run(company)
            company_run.font.size = Pt(11)
            company_run.font.color.rgb = _COLOR_BODY

        if dates:
            dates_para = doc.add_paragraph()
            dates_para.paragraph_format.space_after = Pt(2)
            dates_run = dates_para.add_run(dates)
            dates_run.font.size = Pt(10)
            dates_run.italic = True
            dates_run.font.color.rgb = _COLOR_MUTED

        for bullet in bullets:
            self._add_bullet(doc, bullet)

    def _add_education_entry(self, doc: Document, edu: Dict[str, str]) -> None:
        """Add a single education entry."""
        degree = edu.get("degree", "")
        institution = edu.get("institution", "")
        dates = edu.get("dates", "")

        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(4)
        para.paragraph_format.space_after = Pt(2)

        deg_run = para.add_run(degree)
        deg_run.bold = True
        deg_run.font.size = Pt(11)
        deg_run.font.color.rgb = _COLOR_SUBHEADING

        if institution:
            sep_run = para.add_run("  ·  ")
            sep_run.font.color.rgb = _COLOR_MUTED
            inst_run = para.add_run(institution)
            inst_run.font.size = Pt(11)
            inst_run.font.color.rgb = _COLOR_BODY

        if dates:
            date_para = doc.add_paragraph()
            date_para.paragraph_format.space_after = Pt(2)
            date_run = date_para.add_run(dates)
            date_run.italic = True
            date_run.font.size = Pt(10)
            date_run.font.color.rgb = _COLOR_MUTED

    def _add_skills(self, doc: Document, skills: List[str]) -> None:
        """Add skills as a wrapped comma-separated line (ATS-friendly)."""
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(4)
        run = para.add_run("  ·  ".join(skills))
        run.font.size = Pt(11)
        run.font.color.rgb = _COLOR_BODY
